from docx.shared import Cm
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx import Document

def get_template(filename):
    file = open(filename, 'r')
    logo = file.readline().split(':')[1].strip()

    title = file.readline().split(':')[1].strip()
    content = file.readline().split(':')[1].strip()
    agendas = file.readline().split(':')[1].strip()
    date = file.readline().split(':')[1].strip()
    author = file.readline().split(':')[1].strip()
    
    return logo, title, content, agendas, date, author

# 로고 그림을 추가하는 함수
def add_logo(document, logo_path):
    logo_p = document.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logo_run = logo_p.add_run()
    logo_run.add_picture(logo_path, width=Cm(3), height=Cm(3))
    logo_run.add_break(WD_BREAK.LINE)
    logo_run.add_break(WD_BREAK.LINE)

# 타이틀을 추가하는 함수
def add_title(document, title):
    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font_name = '새굴림'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '새굴림')
    title_run.add_break(WD_BREAK.LINE)

# 본문을 추가하는 함수
def add_content(document, content):
    content_p = document.add_paragraph()
    content_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    content_run = content_p.add_run(content)
    content_run.font.size = Pt(12)
    content_run.font_name = '새굴림'
    content_run._element.rPr.rFonts.set(qn('w:eastAsia'), '새굴림')
    content_run.add_break(WD_BREAK.LINE)

# 안건을 추가하는 함수
def add_agendas(document, agendas):
    for agenda in agendas:
        agenda_p = document.add_paragraph(agenda, style='List Bullet')
        agenda_run = agenda_p.runs[0]
        agenda_run.font.size = Pt(12)
        agenda_run.font_name = '새굴림'
        agenda_run._element.rPr.rFonts.set(qn('w:eastAsia'), '새굴림')

# 날짜를 추가하는 함수
def add_date(document, date):
    date_p = document.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run(date)
    date_run.font.size = Pt(12)
    date_run.font_name = '새굴림'
    date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '새굴림')
    date_run.add_break(WD_BREAK.LINE)

def main():
    logo, title, content, agendas, date, author = get_template('템플릿파일.txt')
    
    filename = '보고서.docx'
    document = Document()

    add_logo(document, logo)
    add_title(document, title)
    add_content(document, content)
    add_agendas(document, agendas.split(','))
    add_date(document, date)
    add_date(document, author)

    document.save(filename)

if __name__ == "__main__":
    main()